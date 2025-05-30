import io
import docx

class DocxParser:
    async def parse(self, file_content: bytes) -> str:
        """
        Parse a DOCX file and extract its text content
        
        Args:
            file_content: Binary content of the DOCX file
            
        Returns:
            Extracted text from the DOCX
        """
        docx_file = io.BytesIO(file_content)
        
        try:
            doc = docx.Document(docx_file)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"
            
            # Add double newline between sections
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text:
                                text += paragraph.text + "\n"
                    text += "\n"
                text += "\n"
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {type(e).__name__}: {str(e)}")